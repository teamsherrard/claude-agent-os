#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_doc.py — Realtor AI Brain universal document renderer.

Turns the system's house-style STRUCTURED TEXT (the format every skill already writes:
CAPS section bands, • bullets, "Label:" lead-ins, simple aligned tables, the credibility
stamp) into a clean, premium .docx — one neutral standard for every client.

  Look:  Arial · pure-black text · light-grey rules · near-black table headers (white text)
         · real bullet lists · real tables · NO client branding, NO colour.

Usage:
  python3 render_doc.py  INPUT.txt  "OUTPUT.docx"
      [--title "Doc Title"]        # overrides the first line as the big title
      [--subtitle "Agent · City"]  # a subtitle line under the title

The skills keep producing the same structured text; only the final SAVE step changes
(render to .docx instead of saving raw text/plain). Requires python-docx.
"""
import sys, re, argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("ERROR: python-docx is required.  Install it with:  pip install python-docx\n")
    sys.exit(2)

# ---------- the one house style (neutral PREMIUM — no client branding, no colour) ----------
INK    = RGBColor(0x11, 0x11, 0x11)  # title + headings — near-black (softer + more premium than pure black)
BODY   = RGBColor(0x1A, 0x1A, 0x1A)  # body copy — near-black, crisp and legible
MUTED  = RGBColor(0x60, 0x63, 0x67)  # refined mid-grey — eyebrow / subtitle / meta / small print
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RULE   = "D7DADD"                     # hairline section rule (light, refined)
ACCENT = "111111"                     # near-black accent hairline (cover)
HEAD_FILL, ALT_FILL, COOL_FILL = "202226", "F5F6F7", "F7F8F9"   # refined charcoal header · light rows · callouts
FONT   = "Arial"                      # one clean sans-serif, installed everywhere

def _font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)

def _track(run, val=0):
    # letter-spacing in twentieths of a point (60 = 3pt) — the premium-typography cue
    if not val: return run
    rpr = run._element.get_or_add_rPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(val)); rpr.append(sp)
    return run

def _run(p, text, size=10.5, color=BODY, bold=False, italic=False, track=0):
    r = p.add_run(text); _font(r); r.font.size = Pt(size)
    r.font.color.rgb = color; r.bold = bold; r.italic = italic
    if track: _track(r, track)
    return r

def _sp(p, before=0, after=6, line=1.16):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    return p

def _border(p, color=RULE, sz=12, space=5):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), str(sz))
    bt.set(qn('w:space'), str(space)); bt.set(qn('w:color'), color)
    b.append(bt); pPr.append(b)

def _cell_bg(cell, hexc):
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexc)
    cell._tc.get_or_add_tcPr().append(shd)

def _cell_borders(cell, color="E4E4E6", sz=4):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for s in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz)); e.set(qn('w:color'), color); b.append(e)
    tcPr.append(b)

def _cell_margins(cell, t=80, b=80, l=130, r=130):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for s, v in (('top', t), ('bottom', b), ('start', l), ('end', r)):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

class Doc:
    def __init__(self):
        self.d = Document()
        st = self.d.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5); st.font.color.rgb = BODY
        s = self.d.sections[0]
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
        self.usable = s.page_width - s.left_margin - s.right_margin

    def heading(self, text):
        p = self.d.add_paragraph(); _sp(p, 22, 8, 1.0)
        _run(p, text.upper(), 12, INK, bold=True, track=70)   # tracked small-caps label reads premium
        _border(p, RULE, 6, 7)                                 # hairline, not a heavy bar

    def subheading(self, text, note=""):
        p = self.d.add_paragraph(); _sp(p, 11, 3)
        _run(p, text.strip(), 10.5, INK, bold=True, track=15)
        if note: _run(p, "   " + note, 9.5, MUTED, italic=True)

    def body(self, text, size=10.5):
        p = self.d.add_paragraph(); _sp(p, 0, 7)
        m = re.match(r'^([A-Z][A-Za-z0-9 /&→\-\(\)]{2,40}?):\s\s?(.*)$', text)
        if m and len(m.group(1)) < 38:
            _run(p, m.group(1) + ":  ", size, INK, bold=True); _run(p, m.group(2), size)
        else:
            _run(p, text, size)

    def bullet(self, text, lead=None):
        p = self.d.add_paragraph(); _sp(p, 0, 4)
        p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.18)
        _run(p, "•  ", 10.5, INK, bold=True)
        if lead: _run(p, lead, 10.5, INK, bold=True)
        _run(p, text, 10.5)

    def numbered(self, n, lead, text):
        p = self.d.add_paragraph(); _sp(p, 4, 5)
        p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.first_line_indent = Inches(-0.4)
        _run(p, f"{n}.  ", 12, INK, bold=True)
        if lead: _run(p, lead + " ", 10.5, INK, bold=True)
        _run(p, text, 10.5)

    def cue(self, text):
        p = self.d.add_paragraph(); _sp(p, 1, 3)
        p.paragraph_format.left_indent = Inches(0.28)
        _run(p, text, 9.5, MUTED, italic=True)

    def callout(self, text, label=None):
        t = self.d.add_table(rows=1, cols=1); t.autofit = False
        c = t.rows[0].cells[0]; c.width = self.usable; c.text = ""
        p = c.paragraphs[0]; _sp(p, 2, 2, 1.12)
        if label: _run(p, label + "  ", 10, INK, bold=True)
        _run(p, text, 10, BODY, italic=True)
        _cell_bg(c, COOL_FILL); _cell_borders(c, "E6E3DC"); _cell_margins(c, 90, 90, 150, 150)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def table(self, headers, rows, widths):
        t = self.d.add_table(rows=1, cols=len(headers)); t.autofit = False; t.allow_autofit = False
        for i, h in enumerate(headers):
            self._cell(t.rows[0].cells[i], h, 9.5, WHITE, bold=True, bg=HEAD_FILL)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells; bg = ALT_FILL if ri % 2 else "FFFFFF"
            for i, val in enumerate(row):
                al = 'center' if (i == 0 and len(val) <= 3) else 'left'
                self._cell(cells[i], val, 9.5, BODY, bg=bg, align=al)
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = w
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def _cell(self, cell, text, size, color, bold=False, align='left', bg=None):
        cell.text = ""; p = cell.paragraphs[0]
        p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        _sp(p, 1, 1, 1.05); _run(p, text, size, color, bold)
        if bg: _cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; _cell_borders(cell); _cell_margins(cell)

    def title_block(self, title, subtitle, meta_lines, eyebrow=None):
        if eyebrow:
            p = self.d.add_paragraph(); _sp(p, 22, 6)
            _run(p, eyebrow.upper(), 9, MUTED, bold=True, track=140)   # letter-spaced kicker
        else:
            self.d.add_paragraph().paragraph_format.space_after = Pt(12)   # top air
        p = self.d.add_paragraph(); _sp(p, 0, 2, 1.02)
        _run(p, title, 30, INK, bold=True, track=-8)                  # larger, tightly tracked display
        if subtitle:
            p = self.d.add_paragraph(); _sp(p, 1, 4)
            _run(p, subtitle, 12.5, MUTED, track=20)                  # muted subtitle = refined, not shouty
        p = self.d.add_paragraph(); _sp(p, 4, 7); _border(p, ACCENT, 8, 7)   # near-black accent hairline
        for ml in meta_lines:
            p = self.d.add_paragraph(); _sp(p, 3, 0)
            if ml.lower().startswith("powered by"):
                _run(p, ml.upper(), 8.5, MUTED, bold=True, track=60)
            else:
                _run(p, ml, 9.5, MUTED)
        self.d.add_paragraph().paragraph_format.space_after = Pt(4)   # breathing room before the body

    def stamp(self, text):
        p = self.d.add_paragraph(); _sp(p, 8, 2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text.upper(), 9, MUTED, bold=True)

    def footer_note(self, text):
        p = self.d.add_paragraph(); _sp(p, 8, 0); _border(p, "D9D9D9", 6, 6)
        cp = self.d.add_paragraph(); _sp(cp, 4, 0); _run(cp, text, 8.5, MUTED, italic=True)

    def save(self, path): self.d.save(path)

# ---------- parsing the house-style structured text ----------
_DIV = set("═─—–-")
def is_band(l):  # a full-width divider rule, any style (═ ─ — – -)
    s = l.strip(); return len(s) > 8 and bool(s) and set(s) <= _DIV
def heavy(l): return bool(l.strip()) and set(l.strip()) == {"═"} and len(l.strip()) > 8
def heading_like(s):  # CAPS-dominant (ignoring a lowercase parenthetical) → header, not a footer sentence
    base = s.split("(")[0]
    letters = [c for c in base if c.isalpha()]
    if not letters: return False
    return sum(c.isupper() for c in letters) / len(letters) >= 0.6 and len(s) <= 80
NUM = re.compile(r'^\s*(\d+)\s')
SUBBAND = re.compile(r'^\s*[─—-]{2,}\s*([A-Z0-9][A-Za-z0-9 &/\-\']{2,40}?)\s*[─—-]{2,}\s*$')  # ── LABEL ──
CUE = re.compile(r'^\s*(>>|\[|FACT:|ON SCREEN|PAUSE)')

def render(lines, doc, title=None, subtitle=None, eyebrow=None):
    n = len(lines); i = 0
    # ----- title block: everything before the first band -----
    head = []
    while i < n and not is_band(lines[i]):
        if lines[i].strip(): head.append(lines[i].strip())
        i += 1
    t = title or (head[0] if head else "Document")
    meta = head[1:]
    sub = subtitle
    if sub is None and meta and "·" in meta[0] and not meta[0].lower().startswith(("prepared", "powered")):
        sub, meta = meta[0], meta[1:]
    doc.title_block(t, sub, meta, eyebrow=eyebrow)

    # ----- sections -----
    def section_heading_at(k):
        # band, then the heading line. Wrapped both sides (closing band present) => definitely
        # a heading; divider only above => require it to look like a header, so a footer
        # sentence under a rule is NOT promoted to a heading.
        if is_band(lines[k]) and k+1 < n:
            cand = lines[k+1].strip()
            if cand and not is_band(lines[k+1]):
                wrapped = (k+2 < n and is_band(lines[k+2]))
                if wrapped or heading_like(cand):
                    return cand, (k+3 if wrapped else k+2)
        return None, None

    while i < n:
        if is_band(lines[i]):
            h, nxt = section_heading_at(i)
            if h is not None:
                doc.heading(h); i = nxt; continue
            i += 1; continue            # stray rule — skip
        raw = lines[i]; line = raw.strip()
        if line == "": i += 1; continue

        # game-plan style tables (safe: only fire on these exact shapes) ----
        if re.match(r'^\s*#\s+EXACT TITLE', raw):
            tcol = raw.find("EXACT TITLE"); icol = raw.find("SEARCH INTENT"); rows = []; i += 1
            while i < n:
                rl = lines[i]
                if rl.strip() == "": i += 1; break
                if not NUM.match(rl): break
                num = rl[:tcol].strip(); rest = rl[tcol:].rstrip()
                gaps = list(re.finditer(r'\s{2,}', rest))
                if gaps:
                    g = gaps[-1]; ttl, intent = rest[:g.start()].strip(), rest[g.end():].strip()
                else:
                    s = max(0, min(icol - tcol, len(rest) - 1))
                    while s > 0 and rest[s] != ' ': s -= 1
                    ttl, intent = rest[:s].strip(), rest[s:].strip()
                rows.append([num, ttl, intent]); i += 1
            doc.table(["#", "Exact video title", "Search intent & lead type"], rows,
                      [Inches(0.4), Inches(4.05), Inches(2.25)])
            continue
        mweek = re.match(r'^\s*Week\s+(\d+)\s+—\s+(.*)\((Pillar[^)]*)\)\s*$', raw)
        if mweek:
            rows = []
            while i < n:
                m = re.match(r'^\s*Week\s+(\d+)\s+—\s+(.*)\((Pillar[^)]*)\)\s*$', lines[i])
                if not m: break
                rows.append(["Week " + m.group(1), m.group(2).strip(), m.group(3)]); i += 1
            doc.table(["Week", "Video to publish", "Pillar"], rows,
                      [Inches(0.75), Inches(4.4), Inches(1.55)])
            continue
        if '....' in raw and re.match(r'^\s*\S.*\.{3,}', raw):
            rows = []
            while i < n and '....' in lines[i]:
                m = re.match(r'^\s*(.+?)\s*\.{3,}\s*(.+?)\s*$', lines[i])
                if not m: break
                rest = m.group(2).strip(); note = ""
                mm = re.match(r'^(.+?)\s{2,}\((.+)\)\s*$', rest)
                if mm: rest, note = mm.group(1).strip(), mm.group(2).strip()
                rows.append([m.group(1).strip(), rest, note]); i += 1
            doc.table(["Metric", "Target", "Why it matters"], rows,
                      [Inches(2.1), Inches(2.25), Inches(2.35)])
            continue

        # sub-band  "──── LABEL ────"
        msb = SUBBAND.match(raw)
        if msb: doc.subheading(msb.group(1)); i += 1; continue
        # bullets
        if line.startswith("•"):
            bt = line[1:].strip()
            mm = re.match(r'^([^:]{3,60}):\s+(.*)$', bt)
            if mm and ("—" in mm.group(1) or len(mm.group(1)) < 45):
                doc.bullet(mm.group(2), lead=mm.group(1) + ":  ")
            else:
                doc.bullet(bt)
            i += 1; continue
        # numbered list item "1.  TEXT"  (gather indented continuations)
        mnum = re.match(r'^(\d+)\.\s+(.*)$', line)
        if mnum and 1 <= int(mnum.group(1)) <= 9:
            txt = mnum.group(2); i += 1
            while i < n and lines[i].strip() and lines[i].startswith("   ") and not re.match(r'^\s*\d+\.\s', lines[i]) and not is_band(lines[i]):
                txt += " " + lines[i].strip(); i += 1
            ml = re.match(r'^([A-Z][A-Z0-9 ,\-–—/&]{3,55}?[\.\-—])\s*(.*)$', txt)
            if ml: doc.numbered(mnum.group(1), ml.group(1), ml.group(2))
            else:  doc.numbered(mnum.group(1), None, txt)
            continue
        # 4-part-structure style label "HOOK (0:00) — text" + example
        mstruct = re.match(r'^([A-Z][A-Z ]{2,22}?)(\s*\([^)]*\))?\s*—\s*(.+)$', line)
        if mstruct and mstruct.group(1).strip() in (
                "HOOK", "PRIMARY CTA", "SECONDARY CTA", "BODY", "CONTENT", "EARLY CTA", "CLOSING CTA",
                "EARLY WARM CTA", "SOFT MID CTA", "MID CTA", "REASSURANCE", "FINAL CTA", "INTRO"):
            label = mstruct.group(1) + (" " + mstruct.group(2) if mstruct.group(2) else "")
            txt = mstruct.group(3); i += 1
            eg = ""
            while i < n and lines[i].strip() and lines[i].startswith("   ") and not is_band(lines[i]):
                eg += (" " if eg else "") + lines[i].strip(); i += 1
            p = doc.d.add_paragraph(); _sp(p, 6, 2)
            _run(p, label, 10.5, INK, bold=True); _run(p, "  —  " + txt, 10.5, BODY)
            if eg: doc.callout(eg.replace("e.g. ", ""), label="Example" if eg.startswith("e.g.") else None)
            continue
        # cue lines (>> ON SCREEN, [PAUSE], FACT:)
        if CUE.match(line): doc.cue(line); i += 1; continue
        # ALL-CAPS sub-label (optionally with a lowercase parenthetical)
        mlbl = re.match(r'^([A-Z][A-Z0-9 &/\-]{5,42}?)(\s*\(.*\))?:?\s*$', line)
        if mlbl and mlbl.group(1).strip().count(' ') >= 1 and not NUM.match(raw):
            doc.subheading(mlbl.group(1).strip(), mlbl.group(2).strip() if mlbl.group(2) else "")
            i += 1; continue
        # the credibility stamp
        if line.lower().startswith("powered by mike sherrard"):
            doc.stamp(line); i += 1; continue
        if line.lower().startswith("compliance"):
            doc.footer_note(line); i += 1; continue
        # default paragraph
        doc.body(line); i += 1

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input"); ap.add_argument("output")
    ap.add_argument("--title", default=None); ap.add_argument("--subtitle", default=None)
    ap.add_argument("--eyebrow", default=None)
    a = ap.parse_args()
    lines = open(a.input, encoding="utf-8").read().split("\n")
    doc = Doc(); render(lines, doc, a.title, a.subtitle, a.eyebrow); doc.save(a.output)
    print("rendered:", a.output)

if __name__ == "__main__":
    main()
