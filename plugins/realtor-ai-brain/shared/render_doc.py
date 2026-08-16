#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_doc_v2.py — Realtor AI Brain universal document renderer, BOOK edition.

Everything render_doc.py does (CAPS section bands, • bullets, "Label:" lead-ins,
sub-bands, pipe tables, cues, the credibility stamp — Arial, near-black,
neutral-premium, no client branding) PLUS a premium long-form "book" mode for
multi-chapter deliverables like the Business Brain Book.

BOOK MODE — how it activates (the trigger)
------------------------------------------
Book mode turns on when EITHER:
  • the input contains a [[TOC]] ... [[/TOC]] block   (the primary, content-driven
    trigger — skills opt in simply by writing a contents block; no CLI change), OR
  • the --book CLI flag is passed                      (for cover + chapter breaks +
    footer page numbers on a book that has no contents page).
Docs with neither render EXACTLY like the original render_doc.py: no cover page
break, no TOC, no footer part, no bookmarks — byte-level structure preserved.

NEW GRAMMAR (book mode)
-----------------------
1. COVER PAGE — the existing title block (eyebrow → big ~29pt title → accent
   hairline → subtitle/byline/meta lines) is rendered as a true cover with top
   air, followed by a page break before the body.

2. CONTENTS block:
       [[TOC]]
       CHAPTER 3 — MARKET INTELLIGENCE :: one-line summary
       LEAD ENGINE :: one-line summary
       [[/TOC]]
   Renders a "CONTENTS" page: each entry is an internal hyperlink
   (w:hyperlink w:anchor=...) to the bookmark placed on the matching body
   heading, with the summary underneath in 9.5pt grey. No page numbers (not
   computable at render time; the links are the navigation). Page break after.
   Entries may be written WITH or WITHOUT the "CHAPTER n —"/"PART X —" prefix —
   matching normalizes both sides (strip prefix, case, punctuation).
   FAIL-SOFT: an entry with no matching heading renders as plain text and emits
       WARNING: TOC entry "<title>" has no matching heading - rendered as plain text
   on stderr (the verify gate greps for WARNING).

3. CHAPTER / PART bands — a CAPS band heading matching
       ^(PART [IVXLC]+|CHAPTER \\d+) <dash> TITLE
   renders a small grey tracked kicker line ("CHAPTER 3") above the title
   (~15.5pt tracked caps + hairline rule) and takes a PAGE BREAK BEFORE it —
   except when it is the first body element (the cover/TOC already broke).
   EVERY band heading (chapter or plain) is bookmarked with a deterministic slug.

4. ">> " CALLOUT — a line starting ">> " renders as a shaded key-insight box
   (single-cell table, #F3F4F6, slightly inset, bold lead when the text starts
   with a "Label:" pattern). Grayscale only. Book mode only — in legacy docs
   ">>" lines keep their original cue rendering, preserving compatibility.

BOOKMARK SLUG RULE (deterministic)
----------------------------------
  norm  = heading text → strip leading "CHAPTER n"/"PART X" + dash → UPPER
          → every non-[A-Z0-9 ] char becomes a space → collapse whitespace
  slug  = "bm_" + norm.lower() with non-alphanumerics collapsed to "_",
          truncated to 32 chars after the prefix; duplicate slugs get _2, _3 …
          in document order.  (≤ 40 chars, starts with a letter — Word-valid.)
  The TOC matches entries to headings by comparing `norm` on both sides.

LEGACY PARITY (hard requirement — 5 other plugins reuse this file)
------------------------------------------------------------------
  A doc with no [[TOC]] block and no --book flag renders BYTE-IDENTICAL to the
  original render_doc.py: charcoal #202226 table headers with white text,
  10.5pt sub-bands, ">> " lines as muted script cues, no cover break, no
  bookmarks, no footer. Book mode adds (never replaces): 12pt sub-bands,
  ">> " insight boxes (script cue heads ON SCREEN/PAUSE/FACT:/[ exempt), and
  outline levels on PART/CHAPTER bands (Google Docs shows a chapter sidebar).
  Hierarchy: title ~29pt · chapter heading 15.5pt tracked caps · sub-band 12pt
  (book) / 10.5pt (legacy) bold · body 10.5pt · TOC summary 9.5pt grey.
  Underline appears ONLY on hyperlinks.

FOOTER PAGE NUMBERS (book mode)
-------------------------------
A centred PAGE field in the section footer (begin/instrText/separate/"1"/end,
so a cached result shows before any field update).

GOOGLE DOCS CONVERSION (.docx uploaded to Drive, opened in Docs)
----------------------------------------------------------------
SURVIVES: bookmarks + internal w:anchor hyperlinks (become Docs bookmarks/links),
page breaks and w:pageBreakBefore, table cell shading, footer PAGE field
(becomes a live Docs page number), bold/size/colour, cell margins.
KNOWN RISKS: letter-tracking (w:spacing) is dropped by Docs (cosmetic only —
headings simply render untracked); the callout table's left inset (w:tblInd)
may be flattened to full width; field chrome is simplified to a plain page
number. All degrade gracefully; nothing breaks navigation or content.

Usage:
  python3 render_doc_v2.py INPUT.txt "OUTPUT.docx"
      [--title "Doc Title"] [--subtitle "Agent · City"] [--eyebrow "KICKER"]
      [--book]        # force book mode without a [[TOC]] block

Requires python-docx.
"""
import sys, re, argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("ERROR: python-docx is required.  Install it with:  pip install python-docx\n")
    sys.exit(2)

# ---------- the one house style (neutral PREMIUM — no client branding, no colour) ----------
INK    = RGBColor(0x11, 0x11, 0x11)  # title + headings — near-black
BODY   = RGBColor(0x1A, 0x1A, 0x1A)  # body copy — near-black, crisp and legible
MUTED  = RGBColor(0x60, 0x63, 0x67)  # refined mid-grey — eyebrow / subtitle / meta / small print
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RULE   = "D7DADD"                     # hairline section rule
ACCENT = "111111"                     # near-black accent hairline (cover)
HEAD_FILL = "202226"                  # table header row — refined charcoal, white text (house style, all modes)
ALT_FILL, COOL_FILL = "F5F6F7", "F7F8F9"
CALLOUT_FILL = "F3F4F6"               # ">> " key-insight box
FONT   = "Arial"

def _font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(a), name)

def _track(run, val=0):
    if not val: return run
    rpr = run._element.get_or_add_rPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(val)); rpr.append(sp)
    return run

def _run(p, text, size=10.5, color=BODY, bold=False, italic=False, track=0):
    r = p.add_run(text); _font(r); r.font.size = Pt(size)
    r.font.color.rgb = color; r.bold = bold; r.italic = italic
    if track: _track(r, track)
    return r

def _sp(p, before=0, after=6, line=1.16):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    return p

def _border(p, color=RULE, sz=12, space=5):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), str(sz))
    bt.set(qn('w:space'), str(space)); bt.set(qn('w:color'), color)
    b.append(bt); pPr.append(b)

def _cell_bg(cell, hexc):
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexc)
    cell._tc.get_or_add_tcPr().append(shd)

def _cell_borders(cell, color="E4E4E6", sz=4):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for s in ('top', 'bottom', 'left', 'right'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz)); e.set(qn('w:color'), color); b.append(e)
    tcPr.append(b)

def _cell_margins(cell, t=80, b=80, l=130, r=130):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for s, v in (('top', t), ('bottom', b), ('start', l), ('end', r)):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

def _hyperlink(p, anchor, text, size=11, color="111111", bold=False):
    """Internal hyperlink run: <w:hyperlink w:anchor="...">. Underlined (the only
    underline in the house style)."""
    h = OxmlElement('w:hyperlink'); h.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rf.set(qn(a), FONT)
    rPr.append(rf)
    if bold: rPr.append(OxmlElement('w:b'))
    for tag in ('w:sz', 'w:szCs'):
        e = OxmlElement(tag); e.set(qn('w:val'), str(int(round(size * 2)))); rPr.append(e)
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); h.append(r)
    p._p.append(h)

class Doc:
    def __init__(self):
        self.book = False           # set by render() when book mode is active
        self.d = Document()
        st = self.d.styles['Normal']; st.font.name = FONT; st.font.size = Pt(10.5); st.font.color.rgb = BODY
        s = self.d.sections[0]
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
        self.usable = s.page_width - s.left_margin - s.right_margin
        self._bmid = 0

    # ---- bookmarks --------------------------------------------------------
    def _bookmark(self, p, name):
        self._bmid += 1
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), str(self._bmid)); start.set(qn('w:name'), name)
        end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), str(self._bmid))
        el = p._p
        pPr = el.find(qn('w:pPr'))
        if pPr is not None: pPr.addnext(start)
        else: el.insert(0, start)
        el.append(end)

    # ---- headings ---------------------------------------------------------
    def heading(self, text, slug=None):
        p = self.d.add_paragraph(); _sp(p, 22, 8, 1.0)
        _run(p, text.upper(), 12, INK, bold=True, track=70)
        _border(p, RULE, 6, 7)
        if slug: self._bookmark(p, slug)

    def chapter_heading(self, kicker, title_text, slug=None, page_break=True):
        """CHAPTER/PART band: small grey tracked kicker, then a larger tracked-caps
        title with the hairline rule. Page-breaks before itself unless it is the
        first body element."""
        p = self.d.add_paragraph(); _sp(p, 0, 5, 1.0)
        if page_break: p.paragraph_format.page_break_before = True
        _run(p, kicker.upper(), 9.5, MUTED, bold=True, track=140)
        hp = self.d.add_paragraph(); _sp(hp, 0, 10, 1.02)
        _run(hp, title_text.upper(), 15.5, INK, bold=True, track=60)
        _border(hp, RULE, 6, 7)
        # outline level (PART=1, CHAPTER=2) — Google Docs maps these to its
        # outline sidebar, so chapter navigation survives even if a viewer
        # drops the TOC's internal links on conversion.
        lvl = OxmlElement('w:outlineLvl')
        lvl.set(qn('w:val'), '0' if kicker.upper().startswith('PART') else '1')
        hp._p.get_or_add_pPr().append(lvl)
        if slug: self._bookmark(hp, slug)

    def subheading(self, text, note=""):
        p = self.d.add_paragraph(); _sp(p, 11, 3)
        _run(p, text.strip(), 12 if self.book else 10.5, INK, bold=True, track=15)
        if note: _run(p, "   " + note, 9.5, MUTED, italic=True)

    # ---- body elements ----------------------------------------------------
    def body(self, text, size=10.5):
        p = self.d.add_paragraph(); _sp(p, 0, 7)
        m = re.match(r'^([A-Z][A-Za-z0-9 /&→\-\(\)]{2,40}?):\s\s?(.*)$', text)
        if m and len(m.group(1)) < 38:
            _run(p, m.group(1) + ":  ", size, INK, bold=True); _run(p, m.group(2), size)
        else:
            _run(p, text, size)

    def bullet(self, text, lead=None):
        p = self.d.add_paragraph(); _sp(p, 0, 4)
        p.paragraph_format.left_indent = Inches(0.28); p.paragraph_format.first_line_indent = Inches(-0.18)
        _run(p, "•  ", 10.5, INK, bold=True)
        if lead: _run(p, lead, 10.5, INK, bold=True)
        _run(p, text, 10.5)

    def numbered(self, n, lead, text):
        p = self.d.add_paragraph(); _sp(p, 4, 5)
        p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.first_line_indent = Inches(-0.4)
        _run(p, f"{n}.  ", 12, INK, bold=True)
        if lead: _run(p, lead + " ", 10.5, INK, bold=True)
        _run(p, text, 10.5)

    def cue(self, text):
        p = self.d.add_paragraph(); _sp(p, 1, 3)
        p.paragraph_format.left_indent = Inches(0.28)
        _run(p, text, 9.5, MUTED, italic=True)

    def callout(self, text, label=None):
        # legacy example/callout box (struct branch) — unchanged from v1
        t = self.d.add_table(rows=1, cols=1); t.autofit = False
        c = t.rows[0].cells[0]; c.width = self.usable; c.text = ""
        p = c.paragraphs[0]; _sp(p, 2, 2, 1.12)
        if label: _run(p, label + "  ", 10, INK, bold=True)
        _run(p, text, 10, BODY, italic=True)
        _cell_bg(c, COOL_FILL); _cell_borders(c, "E6E3DC"); _cell_margins(c, 90, 90, 150, 150)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def insight(self, text, label=None):
        """">> " key-insight callout: shaded single-cell table, slightly inset,
        neutral grayscale, bold lead when a Label: pattern opens the line."""
        t = self.d.add_table(rows=1, cols=1); t.autofit = False
        ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'), '216'); ind.set(qn('w:type'), 'dxa')
        t._tbl.tblPr.append(ind)
        c = t.rows[0].cells[0]; c.width = self.usable - Inches(0.3); c.text = ""
        p = c.paragraphs[0]; _sp(p, 2, 2, 1.16)
        if label: _run(p, label + ":  ", 10.5, INK, bold=True)
        _run(p, text, 10.5, BODY)
        _cell_bg(c, CALLOUT_FILL); _cell_borders(c, "E0E2E5"); _cell_margins(c, 110, 110, 170, 170)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def table(self, headers, rows, widths):
        t = self.d.add_table(rows=1, cols=len(headers)); t.autofit = False; t.allow_autofit = False
        for i, h in enumerate(headers):
            self._cell(t.rows[0].cells[i], h, 9.5, WHITE, bold=True, bg=HEAD_FILL)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells; bg = ALT_FILL if ri % 2 else "FFFFFF"
            for i, val in enumerate(row):
                al = 'center' if (i == 0 and len(val) <= 3) else 'left'
                self._cell(cells[i], val, 9.5, BODY, bg=bg, align=al)
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = w
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def _cell(self, cell, text, size, color, bold=False, align='left', bg=None):
        cell.text = ""; p = cell.paragraphs[0]
        p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
        _sp(p, 1, 1, 1.05); _run(p, text, size, color, bold)
        if bg: _cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; _cell_borders(cell); _cell_margins(cell)

    # ---- title block (legacy) / cover (book) ------------------------------
    def title_block(self, title, subtitle, meta_lines, eyebrow=None):
        if eyebrow:
            p = self.d.add_paragraph(); _sp(p, 22, 6)
            _run(p, eyebrow.upper(), 9, MUTED, bold=True, track=140)
        else:
            self.d.add_paragraph().paragraph_format.space_after = Pt(12)
        p = self.d.add_paragraph(); _sp(p, 0, 2, 1.02)
        _run(p, title, 30, INK, bold=True, track=-8)
        if subtitle:
            p = self.d.add_paragraph(); _sp(p, 1, 4)
            _run(p, subtitle, 12.5, MUTED, track=20)
        p = self.d.add_paragraph(); _sp(p, 4, 7); _border(p, ACCENT, 8, 7)
        for ml in meta_lines:
            p = self.d.add_paragraph(); _sp(p, 3, 0)
            if ml.lower().startswith("powered by"):
                _run(p, ml.upper(), 8.5, MUTED, bold=True, track=60)
            else:
                _run(p, ml, 9.5, MUTED)
        self.d.add_paragraph().paragraph_format.space_after = Pt(4)

    def cover(self, title, subtitle, meta_lines, eyebrow=None):
        """True cover page: top air → eyebrow → big title → accent hairline →
        byline/date lines → PAGE BREAK before the body."""
        air = self.d.add_paragraph()
        air.paragraph_format.space_before = Pt(150); air.paragraph_format.space_after = Pt(0)
        if eyebrow:
            p = self.d.add_paragraph(); _sp(p, 0, 10)
            _run(p, eyebrow.upper(), 10, MUTED, bold=True, track=160)
        p = self.d.add_paragraph(); _sp(p, 0, 6, 1.04)
        _run(p, title, 29, INK, bold=True, track=-8)
        rp = self.d.add_paragraph(); _sp(rp, 2, 10); _border(rp, ACCENT, 8, 7)
        if subtitle:
            p = self.d.add_paragraph(); _sp(p, 2, 5)
            _run(p, subtitle, 13, MUTED, track=20)
        for ml in meta_lines:
            p = self.d.add_paragraph(); _sp(p, 3, 0)
            if ml.lower().startswith("powered by"):
                _run(p, ml.upper(), 8.5, MUTED, bold=True, track=60)
            else:
                _run(p, ml, 10, MUTED)
        self.d.add_page_break()

    # ---- contents page ----------------------------------------------------
    def toc_page(self, entries, norm_map):
        p = self.d.add_paragraph(); _sp(p, 6, 14, 1.0)
        _run(p, "CONTENTS", 15.5, INK, bold=True, track=90)
        _border(p, RULE, 6, 7)
        for etitle, summary in entries:
            slug = norm_map.get(_norm_title(etitle))
            ep = self.d.add_paragraph(); _sp(ep, 7, 1 if summary else 6, 1.1)
            if slug:
                _hyperlink(ep, slug, etitle, size=11)
            else:
                sys.stderr.write('WARNING: TOC entry "%s" has no matching heading - rendered as plain text\n' % etitle)
                _run(ep, etitle, 11, INK)
            if summary:
                sp2 = self.d.add_paragraph(); _sp(sp2, 0, 7, 1.14)
                sp2.paragraph_format.left_indent = Inches(0.18)
                _run(sp2, summary, 9.5, MUTED)
        self.d.add_page_break()

    # ---- footer page number (book mode) ------------------------------------
    def page_footer(self):
        """Centred PAGE field (with a cached "1" result so viewers that never
        update fields still show a number). Converts to a live page number in
        Google Docs."""
        f = self.d.sections[0].footer
        p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.text = ""; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 6, 0, 1.0)
        def styled():
            r = p.add_run(); _font(r); r.font.size = Pt(8.5); r.font.color.rgb = MUTED
            return r
        r = styled(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r._r.append(fc)
        r = styled(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '; r._r.append(it)
        r = styled(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate'); r._r.append(fc)
        r = styled(); t = OxmlElement('w:t'); t.text = '1'; r._r.append(t)
        r = styled(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end'); r._r.append(fc)

    def stamp(self, text):
        p = self.d.add_paragraph(); _sp(p, 8, 2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text.upper(), 9, MUTED, bold=True)

    def footer_note(self, text):
        p = self.d.add_paragraph(); _sp(p, 8, 0); _border(p, "D9D9D9", 6, 6)
        cp = self.d.add_paragraph(); _sp(cp, 4, 0); _run(cp, text, 8.5, MUTED, italic=True)

    def save(self, path): self.d.save(path)

# ---------- parsing the house-style structured text ----------
_DIV = set("═─—–-")
def is_band(l):
    s = l.strip(); return len(s) > 8 and bool(s) and set(s) <= _DIV
def heavy(l): return bool(l.strip()) and set(l.strip()) == {"═"} and len(l.strip()) > 8
def heading_like(s):
    base = s.split("(")[0]
    letters = [c for c in base if c.isalpha()]
    if not letters: return False
    return sum(c.isupper() for c in letters) / len(letters) >= 0.6 and len(s) <= 80
NUM = re.compile(r'^\s*(\d+)\s')
SUBBAND = re.compile(r'^\s*[─═—-]{2,}\s*([^─═]{2,70}?)\s*[─═—-]{2,}\s*$')
CUE = re.compile(r'^\s*(>>|\[|FACT:|ON SCREEN|PAUSE)')
CHAP = re.compile(r'^(PART\s+[IVXLCDM]+|CHAPTER\s+\d+)\s*[—–:\-]\s*(.+)$')
_CHAP_PREFIX = re.compile(r'^\s*(PART\s+[IVXLCDM]+|CHAPTER\s+\d+)\s*(?:[—–:\-]\s*)?')
_LABEL = re.compile(r'^([A-Z][A-Za-z0-9 /&→\-\(\)]{2,40}?):\s+(.*)$')

def _norm_title(s):
    """Normalize a heading or TOC-entry title for matching: strip CHAPTER/PART
    prefix, uppercase, strip punctuation, collapse whitespace."""
    up = s.strip().upper()
    stripped = _CHAP_PREFIX.sub('', up)
    if not stripped.strip(): stripped = up          # "PART I" alone — keep the words
    stripped = re.sub(r'[^A-Z0-9 ]+', ' ', stripped)
    return re.sub(r'\s+', ' ', stripped).strip()

def _assign_slugs(heading_texts):
    """Deterministic slugs in document order; norm_map maps normalized title →
    slug of the FIRST heading bearing it."""
    slugs, norm_map, used = [], {}, {}
    for h in heading_texts:
        norm = _norm_title(h)
        base = re.sub(r'[^a-z0-9]+', '_', norm.lower()).strip('_')[:32] or 'section'
        name = 'bm_' + base
        k = used.get(name, 0) + 1; used[name] = k
        if k > 1: name = f'{name[:36]}_{k}'
        slugs.append(name)
        norm_map.setdefault(norm, name)
    return slugs, norm_map

def _extract_toc(lines):
    """Pull the [[TOC]] block out of the line stream. Returns (lines-without-toc,
    entries [(title, summary)], has_toc)."""
    out, entries, has, in_toc = [], [], False, False
    for l in lines:
        s = l.strip()
        if s == '[[TOC]]': has = True; in_toc = True; continue
        if s == '[[/TOC]]': in_toc = False; continue
        if in_toc:
            if not s: continue
            title, sep, summ = s.partition('::')
            entries.append((title.strip(), summ.strip() if sep else ''))
            continue
        out.append(l)
    return out, entries, has

def _render_body(lines, i, doc, book_mode, collect=None, slug_iter=None):
    """The section/body loop. collect=list → dry pass recording heading texts;
    slug_iter → real pass consuming precomputed slugs (same order)."""
    n = len(lines)
    last_head = ''
    emitted = False        # has any body element been emitted yet?

    def heading_slug(h):
        if collect is not None:
            collect.append(h); return None
        if slug_iter is not None:
            return next(slug_iter, None)
        return None

    def section_heading_at(k):
        if is_band(lines[k]) and k+1 < n:
            cand = lines[k+1].strip()
            if cand and not is_band(lines[k+1]):
                wrapped = (k+2 < n and is_band(lines[k+2]))
                if wrapped or heading_like(cand):
                    return cand, (k+3 if wrapped else k+2)
        return None, None

    while i < n:
        if is_band(lines[i]):
            h, nxt = section_heading_at(i)
            if h is not None:
                slug = heading_slug(h) if book_mode else None
                mch = CHAP.match(h)
                if book_mode and mch:
                    doc.chapter_heading(mch.group(1), mch.group(2).strip(), slug, page_break=emitted)
                else:
                    doc.heading(h, slug=slug)
                emitted = True; last_head = h; i = nxt; continue
            i += 1; continue            # stray rule — skip
        raw = lines[i]; line = raw.strip()
        if line == "": i += 1; continue
        emitted = True

        # game-plan style tables (safe: only fire on these exact shapes) ----
        if re.match(r'^\s*#\s+EXACT TITLE', raw):
            tcol = raw.find("EXACT TITLE"); icol = raw.find("SEARCH INTENT"); rows = []; i += 1
            while i < n:
                rl = lines[i]
                if rl.strip() == "": i += 1; break
                if not NUM.match(rl): break
                num = rl[:tcol].strip(); rest = rl[tcol:].rstrip()
                gaps = list(re.finditer(r'\s{4,}', rest)) or list(re.finditer(r'\s{2,}', rest))
                if gaps:
                    g = gaps[-1]; ttl, intent = rest[:g.start()].strip(), rest[g.end():].strip()
                    ttl = re.sub(r'\s{2,}', ' ', ttl); intent = re.sub(r'\s{2,}', ' ', intent)
                else:
                    s = max(0, min(icol - tcol, len(rest) - 1))
                    while s > 0 and rest[s] != ' ': s -= 1
                    ttl, intent = rest[:s].strip(), rest[s:].strip()
                rows.append([num, ttl, intent]); i += 1
            doc.table(["#", "Exact video title", "Search intent & lead type"], rows,
                      [Inches(0.4), Inches(4.05), Inches(2.25)])
            continue
        WEEKV = r'^\s*Week\s+(\d+)\s*[·:]\s*Video\s*(\d+)\s*—\s*(.*?)\s*\(([^)]*)\)\s*$'
        mweekv = re.match(WEEKV, raw)
        if mweekv:
            rows = []
            while i < n:
                m = re.match(WEEKV, lines[i])
                if not m: break
                rows.append([f"Week {m.group(1)}", f"#{m.group(2)}", m.group(3).strip(), m.group(4).strip()]); i += 1
            doc.table(["Week", "#", "Video to publish", "Type"], rows,
                      [Inches(0.7), Inches(0.4), Inches(3.85), Inches(1.75)])
            continue
        mweek = re.match(r'^\s*Week\s+(\d+)\s+—\s+(.*)\((Pillar[^)]*)\)\s*$', raw)
        if mweek:
            rows = []
            while i < n:
                m = re.match(r'^\s*Week\s+(\d+)\s+—\s+(.*)\((Pillar[^)]*)\)\s*$', lines[i])
                if not m: break
                rows.append(["Week " + m.group(1), m.group(2).strip(), m.group(3)]); i += 1
            doc.table(["Week", "Video to publish", "Pillar"], rows,
                      [Inches(0.75), Inches(4.4), Inches(1.55)])
            continue
        if '....' in raw and re.match(r'^\s*\S.*\.{3,}\s*\S', raw):
            rows = []
            while i < n and '....' in lines[i]:
                m = re.match(r'^\s*(.+?)\s*\.{3,}\s*(.+?)\s*$', lines[i])
                if not m: break
                rest = m.group(2).strip(); note = ""
                mm = re.match(r'^(.+?)\s{2,}\((.+)\)\s*$', rest)
                if mm: rest, note = mm.group(1).strip(), mm.group(2).strip()
                rows.append([m.group(1).strip(), rest, note]); i += 1
            if 'COMPETITOR' in (last_head or '').upper():
                doc.table(["Channel", "Numbers", "Notes"], rows,
                          [Inches(2.1), Inches(2.25), Inches(2.35)])
            else:
                doc.table(["Metric", "Target", "Why it matters"], rows,
                          [Inches(2.1), Inches(2.25), Inches(2.35)])
            continue

        # generic pipe table:  | Head 1 | Head 2 |  /  | --- | --- |  /  | a | b |
        if line.startswith("|") and line.endswith("|") and line.count("|") >= 3:
            rows = []
            while i < n:
                l = lines[i].strip()
                if not (l.startswith("|") and l.endswith("|")): break
                cells = [c.strip() for c in l.strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                rows = [r + [""] * (ncol - len(r)) for r in rows]
                w = int(doc.usable / ncol)
                doc.table(rows[0], rows[1:] if len(rows) > 1 else [], [w] * ncol)
            continue

        # sub-band  "──── LABEL ────"
        msb = SUBBAND.match(raw)
        if msb: doc.subheading(msb.group(1)); last_head = msb.group(1); i += 1; continue
        # bullets
        if line.startswith("•"):
            bt = line[1:].strip()
            mm = re.match(r'^([^:]{3,60}):\s+(.*)$', bt)
            if mm and ("—" in mm.group(1) or len(mm.group(1)) < 45):
                doc.bullet(mm.group(2), lead=mm.group(1) + ":  ")
            else:
                doc.bullet(bt)
            i += 1; continue
        # numbered list item "1.  TEXT"  (gather indented continuations)
        mnum = re.match(r'^(\d+)\.\s+(.*)$', line)
        if mnum and 1 <= int(mnum.group(1)) <= 99:
            txt = mnum.group(2); i += 1
            while i < n and lines[i].strip() and lines[i].startswith("   ") and not re.match(r'^\s*\d+\.\s', lines[i]) and not is_band(lines[i]):
                txt += " " + lines[i].strip(); i += 1
            ml = re.match(r'^([A-Z][A-Z0-9 ,\-–—/&]{3,55}?[\.\-—])\s*(.*)$', txt)
            if ml: doc.numbered(mnum.group(1), ml.group(1), ml.group(2))
            else:  doc.numbered(mnum.group(1), None, txt)
            continue
        # 4-part-structure style label "HOOK (0:00) — text" + example
        mstruct = re.match(r'^([A-Z][A-Z ]{2,22}?)(\s*\([^)]*\))?\s*—\s*(.+)$', line)
        if mstruct and mstruct.group(1).strip() in (
                "HOOK", "PRIMARY CTA", "SECONDARY CTA", "BODY", "CONTENT", "EARLY CTA", "CLOSING CTA",
                "EARLY WARM CTA", "SOFT MID CTA", "MID CTA", "REASSURANCE", "FINAL CTA", "INTRO"):
            label = mstruct.group(1) + (" " + mstruct.group(2) if mstruct.group(2) else "")
            txt = mstruct.group(3); i += 1
            eg = ""
            _SL = re.compile(r'^\s*([A-Z][A-Z ]{2,22}?)(\s*\([^)]*\))?\s*—\s')
            while (i < n and lines[i].strip() and lines[i].startswith("   ")
                   and not is_band(lines[i]) and not _SL.match(lines[i])):
                eg += (" " if eg else "") + lines[i].strip(); i += 1
            p = doc.d.add_paragraph(); _sp(p, 6, 2)
            _run(p, label, 10.5, INK, bold=True); _run(p, "  —  " + txt, 10.5, BODY)
            if eg: doc.callout(eg.replace("e.g. ", ""), label="Example" if eg.startswith("e.g.") else None)
            continue
        # ">> " key-insight callout (BOOK grammar; legacy docs keep the cue look).
        # Script cue heads stay cues even in book mode — a book embedding script
        # fragments must not turn ">> ON SCREEN — ..." into an insight box.
        if book_mode and line.startswith(">> ") and not re.match(r'^(ON SCREEN|PAUSE|FACT:|\[)', line[3:].strip()):
            txt = line[3:].strip()
            mlab = _LABEL.match(txt)
            if mlab and len(mlab.group(1)) < 38:
                doc.insight(mlab.group(2), label=mlab.group(1))
            else:
                doc.insight(txt)
            i += 1; continue
        # cue lines (>> ON SCREEN, [PAUSE], FACT:)
        if CUE.match(line): doc.cue(line); i += 1; continue
        # ALL-CAPS sub-label (optionally with a lowercase parenthetical)
        mlbl = re.match(r'^([A-Z][A-Z0-9 &/\-]{5,42}?)(\s*\(.*\))?:?\s*$', line)
        if mlbl and mlbl.group(1).strip().count(' ') >= 1 and not NUM.match(raw):
            doc.subheading(mlbl.group(1).strip(), mlbl.group(2).strip() if mlbl.group(2) else "")
            last_head = mlbl.group(1).strip()
            i += 1; continue
        # the credibility stamp
        if line.lower().startswith("powered by mike sherrard"):
            doc.stamp(line); i += 1; continue
        if line.lower().startswith("compliance"):
            doc.footer_note(line); i += 1; continue
        # default paragraph
        doc.body(line); i += 1

def render(lines, doc, title=None, subtitle=None, eyebrow=None, book=False):
    lines, toc_entries, has_toc = _extract_toc(lines)
    book_mode = bool(book or has_toc)
    doc.book = book_mode
    n = len(lines); i = 0
    # ----- title block: everything before the first band -----
    head = []
    while i < n and not is_band(lines[i]) and not SUBBAND.match(lines[i]) and len(head) < 8:
        if lines[i].strip(): head.append(lines[i].strip())
        i += 1
    t = title or (head[0] if head else "Document")
    meta = head[1:]
    sub = subtitle
    if sub is None and meta and "·" in meta[0] and not meta[0].lower().startswith(("prepared", "powered")):
        sub, meta = meta[0], meta[1:]

    if not book_mode:
        # legacy path — identical structure to render_doc.py v1
        doc.title_block(t, sub, meta, eyebrow=eyebrow)
        _render_body(lines, i, doc, False)
        return

    # ----- book path: cover → (dry pass for slugs) → contents → body → footer -----
    doc.cover(t, sub, meta, eyebrow=eyebrow)
    collected = []
    _render_body(lines, i, Doc(), True, collect=collected)          # dry: discover headings in emission order
    slugs, norm_map = _assign_slugs(collected)
    if has_toc:
        doc.toc_page(toc_entries, norm_map)
    _render_body(lines, i, doc, True, slug_iter=iter(slugs))
    doc.page_footer()

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input"); ap.add_argument("output")
    ap.add_argument("--title", default=None); ap.add_argument("--subtitle", default=None)
    ap.add_argument("--eyebrow", default=None)
    ap.add_argument("--book", action="store_true",
                    help="Force book mode (cover page break, chapter page breaks, footer page numbers) "
                         "even without a [[TOC]] block. A [[TOC]] block turns book mode on by itself.")
    a = ap.parse_args()
    lines = open(a.input, encoding="utf-8").read().split("\n")
    doc = Doc(); render(lines, doc, a.title, a.subtitle, a.eyebrow, book=a.book); doc.save(a.output)
    print("rendered:", a.output)

if __name__ == "__main__":
    main()
